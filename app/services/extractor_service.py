"""
extractor_service.py — Extrae datos de documentos históricos usando IA local.

Soporta: PDF, DOCX, XLSX, imágenes (logos)
Extrae:  cliente, contacto, email, industria, servicios, precios, contexto
"""

import os
import re
import json
import base64
import requests
from typing import Optional

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL      = "gemma3:4b"


# ══════════════════════════════════════════════════════════════════════
# EXTRACTORES DE TEXTO POR TIPO DE ARCHIVO
# ══════════════════════════════════════════════════════════════════════

# ── Datos conocidos del proveedor para filtrar ───────────────────────
_PROVEEDOR_PERSONAS = [
    'andres barrientos', 'barrientos cisternas', 'andrés barrientos',
]
_PROVEEDOR_DATOS = [
    '76.771.397', 'la capitanía 80', 'la capitania 80',
    '+56 9 4951 2772', '4951 2772', '20470014891',
    'servicio técnico gamer', 'servicio tecnico gamer',
    'cyber-protection.cl', 'cyberprotection.cl',
]

def _limpiar_datos_proveedor(datos: dict) -> dict:
    """
    Post-procesa el resultado de la IA eliminando datos del proveedor
    que hayan sido incorrectamente asignados al cliente.
    """
    if not datos or 'cliente' not in datos:
        return datos

    cliente = datos.get('cliente', {}) or {}

    # Limpiar contact_name si es personal del proveedor
    contact = (cliente.get('contact_name') or '').lower()
    if any(p in contact for p in _PROVEEDOR_PERSONAS):
        cliente['contact_name'] = None

    # Limpiar email/phone/notes si contienen datos del proveedor
    for campo in ['email', 'phone', 'notes']:
        val = (cliente.get(campo) or '').lower()
        if any(d in val for d in _PROVEEDOR_DATOS):
            cliente[campo] = None

    datos['cliente'] = cliente
    return datos


def extraer_imagenes_pdf(ruta: str, output_dir: str) -> list:
    """
    Extrae imágenes pequeñas del PDF (posibles logos).
    Retorna lista de rutas a las imágenes extraídas.
    Solo extrae de las primeras 2 páginas y filtra por tamaño.
    """
    import subprocess, glob, shutil
    os.makedirs(output_dir, exist_ok=True)
    prefix = os.path.join(output_dir, "logo")

    try:
        # Extraer solo páginas 1-2 donde suelen estar los logos
        subprocess.run([
            "pdfimages", "-png", "-f", "1", "-l", "2", ruta, prefix
        ], capture_output=True, timeout=30)
    except Exception:
        return []

    imagenes = []
    for img_path in sorted(glob.glob(f"{prefix}-*.png")):
        try:
            from PIL import Image as PILImage
            img = PILImage.open(img_path)
            w, h = img.size
            kb = os.path.getsize(img_path) / 1024
            # Filtrar: logo típico < 800x400px, > 5KB (no vacío), no cuadrado grande
            if w < 800 and h < 400 and kb > 3 and w > 50 and h > 20:
                imagenes.append({
                    "path": img_path,
                    "width": w, "height": h, "kb": round(kb, 1)
                })
        except Exception:
            pass

    return imagenes


def extraer_texto_pdf(ruta: str) -> str:
    """Extrae texto de un PDF usando pdfplumber."""
    try:
        import pdfplumber
        texto = []
        with pdfplumber.open(ruta) as pdf:
            for page in pdf.pages[:10]:  # máx 10 páginas
                t = page.extract_text()
                if t:
                    texto.append(t)
        return "\n\n".join(texto)[:8000]
    except Exception as e:
        return f"[Error leyendo PDF: {e}]"


def extraer_texto_docx(ruta: str) -> str:
    """Extrae texto de un archivo Word."""
    try:
        import docx
        doc  = docx.Document(ruta)
        texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # También tablas
        for table in doc.tables:
            for row in table.rows:
                texto += "\n" + " | ".join(c.text for c in row.cells if c.text.strip())
        return texto[:8000]
    except Exception as e:
        return f"[Error leyendo DOCX: {e}]"


def extraer_texto_xlsx(ruta: str) -> str:
    """Extrae texto de un Excel."""
    try:
        import openpyxl
        wb    = openpyxl.load_workbook(ruta, read_only=True, data_only=True)
        lineas = []
        for ws in wb.worksheets[:3]:  # máx 3 hojas
            lineas.append(f"=== Hoja: {ws.title} ===")
            for row in ws.iter_rows(max_row=100, values_only=True):
                celdas = [str(c) for c in row if c is not None]
                if celdas:
                    lineas.append(" | ".join(celdas))
        return "\n".join(lineas)[:8000]
    except Exception as e:
        return f"[Error leyendo XLSX: {e}]"


def extraer_texto_txt(ruta: str) -> str:
    """Lee archivo de texto plano."""
    try:
        with open(ruta, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()[:8000]
    except Exception as e:
        return f"[Error leyendo archivo: {e}]"


def extraer_texto(ruta: str, extension: str) -> str:
    """Router principal de extracción según extensión."""
    ext = extension.lower().strip('.')
    if ext == 'pdf':
        return extraer_texto_pdf(ruta)
    elif ext in ('docx', 'doc'):
        return extraer_texto_docx(ruta)
    elif ext in ('xlsx', 'xls'):
        return extraer_texto_xlsx(ruta)
    elif ext in ('txt', 'md', 'csv'):
        return extraer_texto_txt(ruta)
    else:
        return "[Formato no soportado para extracción de texto]"


# ══════════════════════════════════════════════════════════════════════
# ANÁLISIS CON IA (Ollama)
# ══════════════════════════════════════════════════════════════════════

def analizar_con_ia(texto: str, nombre_archivo: str) -> dict:
    """
    Envía el texto extraído a Ollama y pide un JSON estructurado
    con todos los datos del cliente y la cotización.
    """
    prompt = f"""Eres un asistente especializado en extraer información de propuestas comerciales de ciberseguridad.

REGLA CRÍTICA — DOS PARTES EN UNA PROPUESTA:

PROVEEDOR (quien VENDE): Es quien redacta, firma y emite la propuesta.
  - Tiene su propio RUT en las condiciones comerciales
  - Su nombre aparece en el pie de página, firma y membrete
  - Ejemplos: Cyber-Protection, Cyber-Protection.CL, Servicio Técnico Gamer Chile SPA
  - Su equipo: Andres Barrientos, CTO, ejecutivos internos → NO son el cliente

CLIENTE (quien COMPRA): Es a quien va DIRIGIDA la propuesta.
  - Aparece después de "Elaborado para:", "Preparado para:", "Dirigido a:", "Para:"
  - Puede aparecer solo como logo sin texto (company_name = null en ese caso)
  - Ejemplos: EPSA, UNAB, Apprecio, Banco Estado, municipalidades, empresas del rubro

REGLAS ESTRICTAS:
- contact_name debe ser un contacto de la empresa CLIENTE, NUNCA del proveedor
- Si el único nombre mencionado es del proveedor (ej: Andres Barrientos Cisternas, CTO), pon contact_name: null
- El RUT 76.771.397-5 y la dirección "La Capitanía 80" son del PROVEEDOR — ignóralos para el cliente
- El teléfono +56 9 4951 2772 es del proveedor — ignóralo para el cliente

Nombre del archivo: {nombre_archivo}

DOCUMENTO:
{texto}

Responde ÚNICAMENTE con JSON válido, sin texto adicional, sin backticks:
{{
  "cliente": {{
    "company_name": "nombre de la empresa CLIENTE que recibe la propuesta (null si no aparece en texto)",
    "contact_name": "contacto en la empresa cliente, NO del proveedor (null si no se encuentra)",
    "email": "email del cliente (null si no se encuentra o es del proveedor)",
    "phone": "teléfono del cliente (null si no se encuentra o es del proveedor)",
    "industry": "sector del cliente: educación, banca, retail, sector público, etc.",
    "notes": "contexto del proyecto o necesidad del cliente (null si no hay)"
  }},
  "proveedor_detectado": "nombre exacto del proveedor que redactó la propuesta",
  "servicios": [
    {{
      "nombre": "nombre del servicio o assessment ofrecido al cliente",
      "descripcion": "descripción breve del servicio",
      "precio_uf": número o null
    }}
  ],
  "contexto": "resumen del problema o necesidad del CLIENTE en 2-3 oraciones",
  "confianza": "alta|media|baja según qué tan clara es la identificación del cliente",
  "notas_extraccion": "qué información faltó o fue difícil de extraer"
}}

Solo el JSON, nada más."""

    try:
        r = requests.post(
            OLLAMA_URL,
            json={
                "model":  MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {"num_predict": 800, "temperature": 0.1}
            },
            timeout=120
        )
        r.raise_for_status()
        respuesta = r.json().get("response", "").strip()

        # Limpiar posibles backticks o texto extra
        respuesta = re.sub(r'```json\s*', '', respuesta)
        respuesta = re.sub(r'```\s*', '', respuesta)

        # Extraer solo el JSON si hay texto extra
        match = re.search(r'\{.*\}', respuesta, re.DOTALL)
        if match:
            respuesta = match.group(0)

        return json.loads(respuesta)

    except json.JSONDecodeError as e:
        return {
            "error": f"No se pudo parsear la respuesta de la IA: {e}",
            "cliente": None,
            "servicios": [],
            "contexto": None,
            "confianza": "baja",
            "notas_extraccion": "Error de parseo JSON"
        }
    except requests.exceptions.ConnectionError:
        raise RuntimeError(
            "Ollama no está disponible. "
            "Ejecuta 'ollama serve' antes de usar este módulo."
        )
    except Exception as e:
        raise RuntimeError(f"Error con Ollama: {e}")


# ══════════════════════════════════════════════════════════════════════
# DETECCIÓN DE LOGO EN IMAGEN
# ══════════════════════════════════════════════════════════════════════

def es_imagen_logo(ruta: str) -> bool:
    """Verifica si una imagen podría ser un logo (pequeña, cuadrada o rectangular)."""
    try:
        from PIL import Image
        img = Image.open(ruta)
        w, h = img.size
        # Logos típicamente: < 2000px en cualquier dimensión, aspect ratio razonable
        return max(w, h) < 2000 and 0.2 < (w/h) < 5
    except Exception:
        return False


def guardar_logo(ruta_origen: str, nombre_empresa: str, assets_dir: str) -> Optional[str]:
    """
    Copia el logo al directorio de assets con nombre normalizado.
    Retorna la ruta relativa al logo guardado.
    """
    try:
        from PIL import Image
        import shutil

        os.makedirs(assets_dir, exist_ok=True)
        # Normalizar nombre: "Universidad Andrés Bello" → "universidad_andres_bello"
        nombre_norm = re.sub(r'[^\w\s-]', '', nombre_empresa.lower())
        nombre_norm = re.sub(r'[\s-]+', '_', nombre_norm).strip('_')
        ext         = os.path.splitext(ruta_origen)[1].lower() or '.jpg'
        nombre_arch = f"logo_{nombre_norm}{ext}"
        ruta_dest   = os.path.join(assets_dir, nombre_arch)

        # Convertir a RGB si es necesario y guardar como JPG
        img = Image.open(ruta_origen).convert("RGB")
        img.save(ruta_dest, quality=95)

        return ruta_dest
    except Exception as e:
        print(f"  ⚠️  No se pudo guardar logo: {e}")
        return None


# ══════════════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ══════════════════════════════════════════════════════════════════════

def procesar_archivo(
    ruta: str,
    nombre_archivo: str,
    assets_dir: str
) -> dict:
    """
    Procesa un archivo histórico y extrae toda la información disponible.

    Retorna:
    {
        "tipo": "documento" | "imagen",
        "datos": {...},        # resultado de analizar_con_ia
        "logo_path": str|None, # ruta al logo si se detectó
        "texto_extraido": str  # para debug
    }
    """
    extension = nombre_archivo.rsplit('.', 1)[-1] if '.' in nombre_archivo else ''
    es_imagen = extension.lower() in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp')

    resultado = {
        "tipo":           "imagen" if es_imagen else "documento",
        "nombre_archivo": nombre_archivo,
        "datos":          None,
        "logo_path":      None,
        "texto_extraido": ""
    }

    if es_imagen:
        # Las imágenes se tratan como posibles logos
        if es_imagen_logo(ruta):
            resultado["logo_path"]   = ruta
            resultado["logos_extra"] = []
            resultado["datos"] = {
                "cliente":   {"company_name": None},
                "servicios": [],
                "contexto":  None,
                "confianza": "alta",
                "notas_extraccion": "Imagen detectada como posible logo"
            }
        else:
            resultado["datos"] = {
                "cliente":   {"company_name": None},
                "servicios": [],
                "contexto":  None,
                "confianza": "baja",
                "notas_extraccion": "Imagen demasiado grande para ser un logo"
            }
    else:
        # Extraer texto y analizar con IA
        print(f"  📄 Extrayendo texto de {nombre_archivo}...")
        texto = extraer_texto(ruta, extension)
        resultado["texto_extraido"] = texto[:500]

        # Si es PDF, extraer también logos embebidos
        logos_extraidos = []
        if extension.lower() == "pdf":
            print(f"  🖼️  Extrayendo logos embebidos...")
            tmp_logos_dir = ruta + "_logos"
            logos_raw = extraer_imagenes_pdf(ruta, tmp_logos_dir)
            for img_info in logos_raw:
                logos_extraidos.append(img_info["path"])
            if logos_extraidos:
                print(f"  ✅ {len(logos_extraidos)} posible(s) logo(s) detectado(s)")
        resultado["logos_extra"] = logos_extraidos

        if texto.startswith("[Error"):
            resultado["datos"] = {
                "error": texto,
                "cliente": None,
                "servicios": [],
                "contexto": None,
                "confianza": "baja",
                "notas_extraccion": texto
            }
        else:
            print(f"  🤖 Analizando con IA ({len(texto)} chars)...")
            resultado["datos"] = _limpiar_datos_proveedor(
                analizar_con_ia(texto, nombre_archivo)
            )
            print(f"  ✅ Extracción completada (confianza: {resultado['datos'].get('confianza','?')})")

    return resultado
